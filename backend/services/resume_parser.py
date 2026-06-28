"""
Resume parser service.
Supports PDF (pdfplumber) and DOCX (python-docx).
Extracts structured candidate data using regex + NLP patterns.
"""
import re
import os
import json
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from loguru import logger


# ── Text Extraction ────────────────────────────────────────────────────────────

def extract_text_from_pdf(file_path: str) -> str:
    """Extract raw text from PDF using pdfplumber."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts)
    except ImportError:
        logger.warning("pdfplumber not installed; trying PyPDF2 fallback")
        return _extract_text_pypdf2(file_path)
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise


def _extract_text_pypdf2(file_path: str) -> str:
    """Fallback PDF reader."""
    try:
        import PyPDF2
        text_parts = []
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"PyPDF2 fallback failed: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Extract raw text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise


def extract_strings_from_binary(file_path: str) -> str:
    """Extract printable strings from a binary file (e.g. legacy .doc)."""
    try:
        with open(file_path, 'rb') as f:
            data = f.read()
        printable_chars = bytearray()
        for b in data:
            if 32 <= b <= 126 or b in (10, 13, 9):
                printable_chars.append(b)
            else:
                printable_chars.append(32)
        text = printable_chars.decode('ascii', errors='ignore')
        text = re.sub(r'\s+', ' ', text)
        return text
    except Exception as e:
        logger.error(f"Binary text extraction fallback failed: {e}")
        return ""


def extract_text(file_path: str) -> Tuple[str, str]:
    """
    Auto-detect file type and extract text.
    Returns:
        (extracted_text, file_type)
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    try:
        if ext == ".pdf":
            return extract_text_from_pdf(file_path), "pdf"
        elif ext in (".docx", ".doc"):
            try:
                return extract_text_from_docx(file_path), "docx"
            except Exception as docx_err:
                logger.warning(f"DOCX parser failed for {file_path}: {docx_err}. Trying binary reader fallback...")
                return extract_strings_from_binary(file_path), "doc_binary"
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    except Exception as e:
        logger.warning(f"Standard extraction failed: {e}. Attempting binary stream fallback extraction...")
        fallback_text = extract_strings_from_binary(file_path)
        if len(fallback_text.strip()) > 50:
            return fallback_text, "binary_fallback"
        raise ValueError(f"Failed to extract text from {file_path}: {e}")


# ── NLP Extraction ────────────────────────────────────────────────────────────

def extract_email(text: str) -> Optional[str]:
    pattern = r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}'
    matches = re.findall(pattern, text)
    return matches[0] if matches else None


def extract_phone(text: str) -> Optional[str]:
    """Extract candidate phone numbers, supporting Indian (+91, 91, 0) and US formats."""
    # Indian: starts with 6-9, has 10 digits, optional prefix
    indian_pattern = r'(?:\+?91|0)?[-\s]?[6-9]\d{9}\b'
    us_pattern = r'(?:\+?1[-.\s]?)?(?:\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4}'
    
    # Try Indian pattern first
    indian_matches = re.findall(indian_pattern, text)
    if indian_matches:
        return indian_matches[0].strip()
        
    # Try US pattern fallback
    us_matches = re.findall(us_pattern, text)
    for match in us_matches:
        cleaned = re.sub(r'\D', '', match)
        if len(cleaned) >= 10:
            return match.strip()
    return None


def extract_linkedin(text: str) -> Optional[str]:
    pattern = r'(?:linkedin\.com/in/|linkedin\.com/profile/)[a-zA-Z0-9\-_]+'
    matches = re.findall(pattern, text, re.IGNORECASE)
    return matches[0] if matches else None


def extract_github(text: str) -> Optional[str]:
    pattern = r'(?:github\.com/)[a-zA-Z0-9\-_]+'
    matches = re.findall(pattern, text, re.IGNORECASE)
    return matches[0] if matches else None


def extract_name(text: str) -> Optional[str]:
    """Extract candidate name from the first non-empty lines."""
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    email = extract_email(text) or ""
    for line in lines[:5]:
        if (email and email.lower() in line.lower()):
            continue
        if re.search(r'resume|curriculum|vitae|cv\b', line, re.IGNORECASE):
            continue
        if re.search(r'@|http|www\.|\+1|linkedin|github', line, re.IGNORECASE):
            continue
        if len(line.split()) in [2, 3, 4] and line[0].isupper():
            return line
    return lines[0] if lines else None


def extract_section(text: str, section_keywords: List[str], next_sections: List[str]) -> str:
    """Extract a specific section from resume text."""
    pattern = r'(?i)(?:^|\n)\s*(' + '|'.join(section_keywords) + r')\s*[:\n]'
    next_pattern = r'(?i)(?:^|\n)\s*(' + '|'.join(next_sections) + r')\s*[:\n]'

    start_match = re.search(pattern, text)
    if not start_match:
        return ""

    start_idx = start_match.end()
    end_idx = len(text)
    next_match = re.search(next_pattern, text[start_idx:])
    if next_match:
        end_idx = start_idx + next_match.start()

    return text[start_idx:end_idx].strip()


def extract_skills(text: str) -> List[str]:
    """Extract skills from resume text using common skill keywords."""
    tech_skills = {
        "python", "java", "javascript", "typescript", "c++", "c#", "c", "go", "golang",
        "rust", "ruby", "php", "swift", "kotlin", "scala", "r", "matlab", "bash",
        "machine learning", "deep learning", "tensorflow", "pytorch", "keras",
        "scikit-learn", "pandas", "numpy", "scipy", "matplotlib", "seaborn",
        "plotly", "opencv", "nlp", "computer vision", "transformers", "bert",
        "llm", "langchain", "hugging face", "xgboost", "lightgbm", "catboost",
        "react", "next.js", "vue", "angular", "node.js", "express", "django",
        "flask", "fastapi", "spring boot", "rest api", "graphql", "websocket",
        "html", "css", "tailwind", "bootstrap",
        "sql", "mysql", "postgresql", "mongodb", "redis", "elasticsearch",
        "cassandra", "dynamodb", "sqlite", "oracle", "neo4j", "pinecone",
        "aws", "azure", "gcp", "google cloud", "docker", "kubernetes", "terraform",
        "ansible", "jenkins", "gitlab ci", "github actions", "ci/cd", "linux",
        "bash", "devops", "mlops",
        "tableau", "power bi", "looker", "excel", "google analytics",
        "airflow", "spark", "hadoop", "kafka", "dbt", "snowflake", "bigquery",
        "git", "jira", "confluence", "figma", "postman", "jupyter", "vscode",
        "agile", "scrum", "kanban", "product management", "team leadership",
        "cybersecurity", "penetration testing", "siem", "network security",
        "arduino", "raspberry pi", "esp32", "rtos", "freertos", "embedded c",
        "verilog", "vhdl", "fpga", "i2c", "spi", "uart", "can bus",
    }

    found_skills = set()
    text_lower = text.lower()

    for skill in tech_skills:
        if skill in text_lower:
            found_skills.add(skill.title() if len(skill) > 3 else skill.upper())

    return sorted(list(found_skills))


def extract_education(text: str) -> List[Dict]:
    """Extract education entries."""
    education = []
    degree_keywords = [
        "b.tech", "b.e.", "m.tech", "m.e.", "mba", "bca", "mca", "bsc", "msc",
        "bachelor", "master", "phd", "ph.d", "b.s.", "m.s.", "associate",
        "diploma", "certification", "certificate",
    ]

    edu_section = extract_section(
        text,
        ["education", "academic background", "qualifications"],
        ["experience", "work", "skills", "projects", "certifications", "achievements"],
    )

    if not edu_section:
        return education

    for line in edu_section.split('\n'):
        line = line.strip()
        if not line:
            continue
        line_lower = line.lower()
        if any(kw in line_lower for kw in degree_keywords):
            education.append({
                "raw": line,
                "degree": line,
                "institution": "",
                "year": _extract_year(line),
                "gpa": _extract_gpa(line),
            })

    return education


def extract_experience(text: str) -> List[Dict]:
    """Extract work experience entries."""
    exp_section = extract_section(
        text,
        ["experience", "work experience", "employment", "professional experience"],
        ["education", "skills", "projects", "certifications", "achievements"],
    )

    experiences = []
    if not exp_section:
        return experiences

    date_pattern = r'(?:\d{4}|\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\w*\s+\d{4})'
    blocks = re.split(r'\n(?=\S)', exp_section)

    for block in blocks:
        block = block.strip()
        if len(block) > 10:
            years = re.findall(date_pattern, block, re.IGNORECASE)
            experiences.append({
                "raw": block[:300],
                "duration": " - ".join(years[:2]) if years else "",
                "company": "",
                "role": "",
            })

    return experiences[:10]


def extract_projects(text: str) -> List[Dict]:
    """Extract project entries."""
    proj_section = extract_section(
        text,
        ["projects", "personal projects", "academic projects", "portfolio"],
        ["experience", "education", "skills", "certifications", "achievements"],
    )

    projects = []
    if not proj_section:
        return projects

    blocks = re.split(r'\n(?=[\•\-\*]|\d+\.|\s{0,2}[A-Z])', proj_section)

    for block in blocks:
        block = block.strip().lstrip('•-* ')
        if len(block) > 15:
            lines = block.split('\n')
            projects.append({
                "title": lines[0].strip()[:100],
                "description": '\n'.join(lines[1:])[:500].strip(),
                "technologies": extract_skills(block),
            })

    return projects[:10]


def extract_certifications(text: str) -> List[str]:
    """Extract certifications."""
    cert_section = extract_section(
        text,
        ["certifications", "certificates", "licenses", "credentials"],
        ["experience", "education", "skills", "projects", "achievements"],
    )

    if not cert_section:
        return []

    certs = []
    for line in cert_section.split('\n'):
        line = line.strip().lstrip('•-* ')
        if len(line) > 5:
            certs.append(line)

    return certs[:15]


def _extract_year(text: str) -> Optional[str]:
    years = re.findall(r'\b(?:19|20)\d{2}\b', text)
    return years[-1] if years else None


def _extract_gpa(text: str) -> Optional[str]:
    """Extract CGPA (10.0 scale, 4.0 scale) and percentages (%)."""
    # 1. Match percentages (e.g. 85.5% or 78 %)
    pct_matches = re.findall(r'\b(\d{2}(?:\.\d{1,2})?)\s*%', text)
    if pct_matches:
        return f"{pct_matches[0]}%"
        
    # 2. Match CGPA out of 10 explicitly (e.g. 8.5/10, 9.2 CGPA)
    out_of_10 = re.findall(r'\b(\d{1,2}\.\d{1,2})\s*(?:/|out of)\s*10\b', text, re.IGNORECASE)
    if out_of_10:
        return f"{out_of_10[0]}/10"
        
    # 3. Match CGPA keyword pointer (e.g. CGPA: 8.5)
    cgpa_keyword = re.findall(r'(?:cgpa|gpa|points|pointer)\s*(?:of|:|-)?\s*(\d{1,2}\.\d{1,2})', text, re.IGNORECASE)
    if cgpa_keyword:
        val = float(cgpa_keyword[0])
        if 4.0 < val <= 10.0:
            return f"{val}/10"
        elif val <= 4.0:
            return f"{val}/4"

    # 4. Out of 4.0 explicitly
    out_of_4 = re.findall(r'\b(\d\.\d{1,2})\s*/\s*4\b', text)
    if out_of_4:
        return f"{out_of_4[0]}/4"
        
    # 5. Fallback float match
    gpa_matches = re.findall(r'\b(\d\.\d{1,2})\b', text)
    for g in gpa_matches:
        val = float(g)
        if 4.0 < val <= 10.0:
            return f"{val}/10"
        elif 0.0 <= val <= 4.0:
            return f"{val}/4"
            
    return None


def _extract_backlogs(text: str) -> int:
    """Detect number of active or history backlogs/KTs (Kept Terms) in resume text."""
    no_backlog_pattern = r'\b(?:no|zero|0|nil|nil active)\s*(?:active\s*)?(?:backlogs?|kt|active kt|kept\s*terms?)\b'
    if re.search(no_backlog_pattern, text, re.IGNORECASE):
        return 0
        
    backlog_count_pattern = r'\b(\d+)\s*(?:active\s*)?(?:backlogs?|kt|kept terms?)\b'
    matches = re.findall(backlog_count_pattern, text, re.IGNORECASE)
    if matches:
        return int(matches[0])
        
    if re.search(r'\b(?:backlogs?|kt|kept terms?)\b', text, re.IGNORECASE):
        return 1
        
    return 0


def _detect_indian_education_metadata(text: str) -> dict:
    """Detects Indian education context such as boards and universities (AKTU, VTU, CBSE, etc.)."""
    text_lower = text.lower()
    boards = []
    if "cbse" in text_lower:
        boards.append("CBSE")
    if "isc" in text_lower or "icse" in text_lower:
        boards.append("ICSE/ISC")
        
    universities = []
    for uni in ["aktu", "vtu", "jntu", "anna university", "delhi university", "mumbai university", "pune university"]:
        if uni in text_lower:
            universities.append(uni.upper())
            
    return {
        "boards": boards,
        "universities": universities,
        "is_indian_context": bool(boards or universities or "india" in text_lower)
    }


# ── Main Parser ────────────────────────────────────────────────────────────────

def parse_resume(file_path: str) -> Dict:
    """
    Full resume parsing pipeline.
    """
    logger.info(f"Parsing resume: {file_path}")

    raw_text, file_type = extract_text(file_path)

    if not raw_text or len(raw_text.strip()) < 50:
        raise ValueError("Could not extract meaningful text from resume. Please check the file.")

    profile = {
        "name": extract_name(raw_text),
        "email": extract_email(raw_text),
        "phone": extract_phone(raw_text),
        "linkedin": extract_linkedin(raw_text),
        "github": extract_github(raw_text),
        "location": _extract_location(raw_text),
        "skills": extract_skills(raw_text),
        "education": extract_education(raw_text),
        "experience": extract_experience(raw_text),
        "projects": extract_projects(raw_text),
        "certifications": extract_certifications(raw_text),
        "achievements": _extract_achievements(raw_text),
        "summary": _extract_summary(raw_text),
        "raw_text": raw_text,
        "file_type": file_type,
        "total_words": len(raw_text.split()),
        "parsing_confidence": _calculate_confidence(raw_text),
        "backlogs": _extract_backlogs(raw_text),
        "indian_education": _detect_indian_education_metadata(raw_text),
    }

    logger.info(
        f"Resume parsed: {profile['name']} | "
        f"{len(profile['skills'])} skills | "
        f"{len(profile['experience'])} experiences | "
        f"{len(profile['projects'])} projects"
    )

    return profile


def _extract_location(text: str) -> Optional[str]:
    location_pattern = r'\b([A-Z][a-z]+(?:[\s,]+[A-Z][a-zA-Z]+)*,\s*(?:[A-Z]{2}|[A-Za-z]+))\b'
    matches = re.findall(location_pattern, text)
    if matches:
        return matches[0]
    return None


def _extract_summary(text: str) -> str:
    summary = extract_section(
        text,
        ["summary", "objective", "profile", "about", "professional summary"],
        ["education", "experience", "skills", "projects"],
    )
    return summary[:500] if summary else ""


def _extract_achievements(text: str) -> List[str]:
    ach_section = extract_section(
        text,
        ["achievements", "accomplishments", "awards", "honors"],
        ["education", "experience", "skills", "projects", "certifications"],
    )

    if not ach_section:
        return []

    achievements = []
    for line in ach_section.split('\n'):
        line = line.strip().lstrip('•-* ')
        if len(line) > 10:
            achievements.append(line)
    return achievements[:10]


def _calculate_confidence(text: str) -> str:
    score = 0
    if extract_email(text):
        score += 20
    if extract_phone(text):
        score += 15
    if extract_skills(text):
        score += 25
    if len(text.split()) > 200:
        score += 20
    if len(text.split()) > 500:
        score += 20

    if score >= 80:
        return "High"
    elif score >= 50:
        return "Medium"
    else:
        return "Low"
